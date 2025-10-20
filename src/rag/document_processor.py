"""
Document processor for extracting text from various file formats
"""
import io
import json
from typing import List, Dict, Optional, Any, Tuple
import PyPDF2
import docx
import markdown
import re
from loguru import logger
from config.settings import settings


class DocumentProcessor:
    """Process documents and extract text content"""
    
    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        # For coordinate tracking
        self.line_endings = []  # Track line break positions
        self.page_breaks = []   # Track page break positions
    
    def extract_text(self, content: bytes, filename: str, content_type: str = None) -> Optional[str]:
        """Extract text from document content based on file type"""
        try:
            # Reset coordinate tracking for new document
            self.line_endings = []
            self.page_breaks = []
            
            # Determine file type from filename if content_type not provided
            if not content_type:
                content_type = self._get_content_type_from_filename(filename)
            
            if content_type == 'application/pdf' or filename.lower().endswith('.pdf'):
                return self._extract_from_pdf_with_coordinates(content)
            elif content_type.startswith('application/vnd.openxmlformats') or filename.lower().endswith('.docx'):
                return self._extract_from_docx_with_coordinates(content)
            elif content_type == 'text/markdown' or filename.lower().endswith(('.md', '.markdown')):
                return self._extract_from_markdown_with_coordinates(content)
            elif content_type == 'application/json' or filename.lower().endswith('.json'):
                return self._extract_from_json_with_coordinates(content)
            elif content_type.startswith('text/') or filename.lower().endswith('.txt'):
                return self._extract_from_text_with_coordinates(content)
            else:
                logger.warning(f"Unsupported file type for {filename}: {content_type}")
                return None
                
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}")
            return None
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF content"""
        text = ""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
        return text
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            doc = docx.Document(io.BytesIO(content))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    def _extract_from_markdown(self, content: bytes) -> str:
        """Extract text from Markdown content"""
        try:
            text = content.decode('utf-8')
            # Convert markdown to plain text (remove markdown formatting)
            html = markdown.markdown(text)
            # Simple HTML tag removal (for basic cases)
            import re
            clean_text = re.sub('<.*?>', '', html)
            return clean_text
        except Exception as e:
            logger.error(f"Error extracting Markdown text: {e}")
            raise
    
    def _extract_from_json(self, content: bytes) -> str:
        """Extract text from JSON content by converting structured data to readable text"""
        try:
            text = content.decode('utf-8')
            json_data = json.loads(text)
            
            # Convert JSON to readable text format
            readable_text = self._json_to_text(json_data)
            return readable_text
            
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing JSON: {e}")
            raise
        except Exception as e:
            logger.error(f"Error extracting JSON text: {e}")
            raise
    
    def _extract_from_text(self, content: bytes) -> str:
        """Extract text from plain text content"""
        try:
            return content.decode('utf-8')
        except Exception as e:
            logger.error(f"Error extracting plain text: {e}")
            raise
    
    def _get_content_type_from_filename(self, filename: str) -> str:
        """Determine content type from filename extension"""
        extension = filename.lower().split('.')[-1]
        content_type_map = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'txt': 'text/plain',
            'md': 'text/markdown',
            'markdown': 'text/markdown',
            'json': 'application/json',
        }
        return content_type_map.get(extension, 'application/octet-stream')
    
    def chunk_text(self, text: str, metadata: Dict = None) -> List[Dict]:
        """Split text into chunks with overlap"""
        if not text or not text.strip():
            return []
        
        chunks = []
        
        # Split by sentences first, then by words if sentences are too long
        sentences = self._split_into_sentences(text)
        
        current_chunk = ""
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence.split())
            
            # If adding this sentence would exceed chunk size, save current chunk
            if current_size + sentence_size > self.chunk_size and current_chunk:
                chunks.append(self._create_chunk(current_chunk, metadata, len(chunks)))
                
                # Start new chunk with overlap from previous chunk
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sentence if overlap_text else sentence
                current_size = len(current_chunk.split())
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_size += sentence_size
        
        # Add the last chunk if it has content
        if current_chunk.strip():
            chunks.append(self._create_chunk(current_chunk, metadata, len(chunks)))
        
        logger.info(f"Created {len(chunks)} chunks from text")
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        import re
        # Simple sentence splitting - can be improved with NLTK or spaCy
        sentences = re.split(r'[.!?]+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of current chunk"""
        words = text.split()
        overlap_words = words[-self.chunk_overlap:] if len(words) > self.chunk_overlap else words
        return " ".join(overlap_words)
    
    def _create_chunk(self, text: str, metadata: Dict, chunk_index: int) -> Dict:
        """Create a chunk dictionary with metadata"""
        return {
            'text': text.strip(),
            'chunk_index': chunk_index,
            'word_count': len(text.split()),
            'char_count': len(text),
            'metadata': metadata or {}
        }
    
    def _json_to_text(self, obj: Any, prefix: str = "", level: int = 0) -> str:
        """Convert JSON object to readable text format"""
        result = []
        indent = "  " * level
        
        if isinstance(obj, dict):
            for key, value in obj.items():
                if isinstance(value, (dict, list)):
                    result.append(f"{indent}{prefix}{key}:")
                    result.append(self._json_to_text(value, "", level + 1))
                else:
                    # Format key-value pairs for better readability
                    formatted_value = self._format_json_value(value)
                    result.append(f"{indent}{prefix}{key}: {formatted_value}")
                    
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                if isinstance(item, (dict, list)):
                    result.append(f"{indent}Item {i + 1}:")
                    result.append(self._json_to_text(item, "", level + 1))
                else:
                    formatted_value = self._format_json_value(item)
                    result.append(f"{indent}- {formatted_value}")
                    
        else:
            return self._format_json_value(obj)
        
        return "\n".join(result)
    
    def _format_json_value(self, value: Any) -> str:
        """Format individual JSON values for better readability"""
        if isinstance(value, str):
            return value
        elif isinstance(value, bool):
            return "Yes" if value else "No"
        elif isinstance(value, (int, float)):
            return str(value)
    
    def _extract_from_pdf_with_coordinates(self, content: bytes) -> str:
        """Extract text from PDF content with coordinate tracking"""
        text = ""
        char_position = 0
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page_num, page in enumerate(pdf_reader.pages):
                # Mark page break position
                if page_num > 0:
                    self.page_breaks.append(char_position)
                
                page_text = page.extract_text()
                page_header = f"\n--- Page {page_num + 1} ---\n"
                
                # Track line endings in the page header
                for i, char in enumerate(page_header):
                    if char == '\n':
                        self.line_endings.append(char_position + i)
                
                char_position += len(page_header)
                text += page_header
                
                # Track line endings in the page content
                for i, char in enumerate(page_text):
                    if char == '\n':
                        self.line_endings.append(char_position + i)
                
                char_position += len(page_text) + 1  # +1 for added newline
                text += page_text + "\n"
                
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
        
        return text
    
    def _extract_from_docx_with_coordinates(self, content: bytes) -> str:
        """Extract text from DOCX content with coordinate tracking"""
        try:
            doc = docx.Document(io.BytesIO(content))
            text = ""
            char_position = 0
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text + "\n"
                
                # Track line ending for this paragraph
                self.line_endings.append(char_position + len(paragraph.text))
                
                text += para_text
                char_position += len(para_text)
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    def _extract_from_markdown_with_coordinates(self, content: bytes) -> str:
        """Extract text from Markdown content with coordinate tracking"""
        try:
            text = content.decode('utf-8')
            char_position = 0
            
            # Track line endings in original text
            for i, char in enumerate(text):
                if char == '\n':
                    self.line_endings.append(i)
            
            # Convert markdown to plain text (remove markdown formatting)
            html = markdown.markdown(text)
            # Simple HTML tag removal (for basic cases)
            clean_text = re.sub('<.*?>', '', html)
            
            return clean_text
            
        except Exception as e:
            logger.error(f"Error extracting Markdown text: {e}")
            raise
    
    def _extract_from_json_with_coordinates(self, content: bytes) -> str:
        """Extract text from JSON content with coordinate tracking"""
        try:
            text = content.decode('utf-8')
            json_data = json.loads(text)
            
            # Convert JSON to readable text format
            readable_text = self._json_to_text(json_data)
            
            # Track line endings in the readable text
            char_position = 0
            for char in readable_text:
                if char == '\n':
                    self.line_endings.append(char_position)
                char_position += 1
            
            return readable_text
            
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing JSON: {e}")
            raise
        except Exception as e:
            logger.error(f"Error extracting JSON text: {e}")
            raise
    
    def _extract_from_text_with_coordinates(self, content: bytes) -> str:
        """Extract text from plain text content with coordinate tracking"""
        try:
            text = content.decode('utf-8')
            
            # Track line endings
            for i, char in enumerate(text):
                if char == '\n':
                    self.line_endings.append(i)
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting plain text: {e}")
            raise
    
    def get_text_coordinates(self, char_start: int, char_end: int) -> Dict[str, Any]:
        """Get coordinate information for a specific text range"""
        try:
            # Find line numbers for start and end positions
            start_line = self._get_line_number(char_start)
            end_line = self._get_line_number(char_end)
            
            # Find page numbers if available
            start_page = self._get_page_number(char_start)
            end_page = self._get_page_number(char_end)
            
            return {
                'char_start': char_start,
                'char_end': char_end,
                'start_line': start_line,
                'end_line': end_line,
                'start_page': start_page,
                'end_page': end_page,
                'line_span': end_line - start_line + 1 if start_line and end_line else 1
            }
            
        except Exception as e:
            logger.error(f"Error getting text coordinates: {e}")
            return {'char_start': char_start, 'char_end': char_end}
    
    def _get_line_number(self, char_position: int) -> Optional[int]:
        """Get line number for a character position"""
        if not self.line_endings:
            return None
        
        line_number = 1
        for line_end in self.line_endings:
            if char_position <= line_end:
                return line_number
            line_number += 1
        
        return line_number
    
    def _get_page_number(self, char_position: int) -> Optional[int]:
        """Get page number for a character position"""
        if not self.page_breaks:
            return 1  # Single page document
        
        page_number = 1
        for page_break in self.page_breaks:
            if char_position < page_break:
                return page_number
            page_number += 1
        
        return page_number
    
    def find_text_coordinates(self, full_text: str, search_text: str, context_chars: int = 50) -> List[Dict[str, Any]]:
        """Find coordinates for specific text within the full document text"""
        results = []
        search_text = search_text.strip()
        
        if not search_text:
            return results
        
        # Find all occurrences of the search text
        start_pos = 0
        while True:
            pos = full_text.find(search_text, start_pos)
            if pos == -1:
                break
            
            # Get coordinates for this occurrence
            coordinates = self.get_text_coordinates(pos, pos + len(search_text))
            
            # Add context around the found text
            context_start = max(0, pos - context_chars)
            context_end = min(len(full_text), pos + len(search_text) + context_chars)
            context = full_text[context_start:context_end]
            
            coordinates.update({
                'found_text': search_text,
                'context': context,
                'context_start': context_start,
                'context_end': context_end
            })
            
            results.append(coordinates)
            start_pos = pos + 1  # Continue searching after this occurrence
        
        return results
